# -*- coding: utf-8 -*-
"""PaperMind 论文精读平台 — 数据库设计文档生成器"""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cn(run, size=None, bold=None, color=None):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, 16, True, RGBColor(0x1E, 0x29, 0x3B))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, 13, True, RGBColor(0x3B, 0x82, 0xF6))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn(r, 11.5, True, RGBColor(0x33, 0x41, 0x55))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p

def para(text, size=10.5, bold=False, color=None, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_cn(r, size, bold, color)
    p.paragraph_format.space_after = Pt(3)
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def table(headers, rows, widths=None, head_bg='1E293B'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        set_cn(r, 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
        shade(c, head_bg)
        if widths: c.width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(v))
            set_cn(r, 9)
            if widths: cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    for line in text.split('\n'):
        r = p.add_run(line + '\n')
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x0F, 0x34, 0x66)

# ================= 封面 =================
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PaperMind 论文精读平台'); set_cn(r, 26, True, RGBColor(0x1E, 0x29, 0x3B))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('数据库设计文档'); set_cn(r, 18, True, RGBColor(0x3B, 0x82, 0xF6))
doc.add_paragraph()
for line in ['文档版本：V2.0（全功能定稿）', '技术栈：MySQL 8.0 + Django ORM + Vue3', '覆盖范围：用户体系 / 文献库（回收站） / 收藏网站（arXiv 真实检索） / 导入文献 / 阅读器（划词·截图OCR·多格式·大纲跳转）/', '　　　　　　知识图谱 / 知识库 / 社区交流 / 团队 / 论文分享 / AI 对话 / Obsidian 导出 / 引擎配置 / 数据备份', '生成日期：2026-08-16']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); set_cn(r, 11, False, RGBColor(0x64, 0x74, 0x8B))
doc.add_page_break()

# ================= 1 概述 =================
h1('1. 文档概述')
para('本文档为 PaperMind 论文精读平台（原型已实现全部交互功能）的数据库设计。平台核心价值：AI 辅助英文论文精读 —— 双语对照、划词翻译/总结、截图 OCR 识别与 AI 翻译总结、多格式文件阅读（PDF/Word/PPT）、个人文献库、收藏网站、知识图谱与知识库沉淀（术语/笔记/卡片/Obsidian 导出），以及团队协作：社区交流（公开笔记动态流/论文社区聚合）、团队（创建/申请/审批）、论文分享（个人/好友/团队）、笔记可见性（公开/仅好友/私密）。')
para('设计原则：', bold=True)
for t in ['① 用户级数据隔离：除论文基础元数据外，所有业务数据（收藏网站、笔记、术语、图谱、偏好）均按 user_id 归属个人；',
          '② JSON 字段承载配置类数据（引擎配置），避免过度建表；',
          '③ 状态用业务枚举字符串（可读、可扩展），不引入独立字典表；',
          '④ 个人文献库与图谱共享论文实体，图谱外部论文（related 类型）与库内论文通过 paper_id 关联。']:
    para(t, indent=0.4)

h2('1.1 功能模块 → 数据表映射')
table(['功能模块', '涉及数据表'], [
    ['用户注册/登录/个人资料（角色/研究方向/头像）', 'users, user_preferences'],
    ['文献库（我的文献：类别/状态/标签/星级/进度）', 'papers, categories, reading_records'],
    ['回收站（删除→恢复→彻底删除，软删除）', 'papers.deleted_at, paper_notes.deleted_at'],
    ['收藏网站（默认预置 6 个/自定义 + arXiv 真实检索）', 'user_sources'],
    ['导入文献（本地文件单个/批量）', 'import_records, papers, paper_files'],
    ['阅读器多格式（PDF/Word/PPT）+ 章节大纲跳转', 'paper_files, papers, paper_notes'],
    ['划词翻译/总结/高亮/笔记（可人工修改 + 三档可见性）', 'paper_highlights, paper_notes'],
    ['截图 OCR + AI 翻译总结（可人工修改）', 'ocr_records, paper_notes'],
    ['AI 对话 / 问 AI', 'ai_conversations, ai_messages'],
    ['知识图谱（节点/边/点击推荐）', 'graph_nodes, graph_edges, papers'],
    ['知识库（卡片/术语/笔记聚合/统计卡跳转）', 'papers, glossary_terms, paper_notes, reading_records'],
    ['导出 Obsidian（逐篇 md）', 'export_records, papers'],
    ['数据备份（个人中心导出/导入 JSON）', 'export_records（backup_type 扩展）'],
    ['偏好设置（翻译/OCR 引擎、界面语言、测试连接）', 'user_preferences'],
    ['社区交流（动态流/论文社区聚合/点赞评论）', 'paper_notes(visibility=public), friendships'],
    ['团队（创建/申请/审批/成员管理）', 'teams, team_members, team_applications'],
    ['论文分享（个人/好友/团队 + 分享管理）', 'paper_shares, friendships'],
], widths=[5.5, 10.5])

# ================= 2 约定 =================
h1('2. 设计约定')
table(['约定项', '规则'], [
    ['字符集', 'utf8mb4 / utf8mb4_unicode_ci（支持 emoji 头像 🦉、中文、公式符号）'],
    ['主键', 'BIGINT UNSIGNED 自增，命名 id'],
    ['外键', '逻辑外键 + 物理索引（字段名 *_id）；物理外键约束由 Django ORM 层控制，便于分库分表'],
    ['时间字段', 'created_at DATETIME DEFAULT CURRENT_TIMESTAMP；updated_at DATETIME ON UPDATE CURRENT_TIMESTAMP'],
    ['软删除', '核心业务表（papers/notes）预留 deleted_at DATETIME NULL'],
    ['JSON 字段', 'MySQL 8.0 JSON 类型，用于引擎配置/批量导出清单'],
    ['枚举取值', '业务字符串：status(想读/在读/读完)、role(普通用户/技术负责人/团队管理员/专业版)、node_type(paper/concept/related)、source(selection/ocr)'],
    ['引擎/Key 安全', '第三方 API Key 存入 user_preferences 的 JSON，正式环境建议加密存储（Django Fernet）'],
], widths=[3.2, 12.8])

# ================= 3 ER =================
h1('3. ER 关系总览')
para('核心关系（1—N 用 ├─，1—1 用 ═）：', size=10)
code_block('''users
 ├─ user_preferences      (1—1 偏好/引擎配置/语言)
 ├─ categories            (类别/文件夹)
 ├─ papers                (我的文献)
 │    ├─ paper_files      (多格式文件 PDF/Word/PPT)
 │    ├─ paper_highlights (划词高亮)
 │    ├─ paper_notes      (划词/截图OCR 笔记，含 AI 翻译总结)
 │    ├─ ocr_records      (截图 OCR 记录)
 │    └─ reading_records  (阅读进度/时长)
 ├─ user_sources          (收藏网站，新用户预置 6 个默认)
 ├─ glossary_terms        (术语库，跨论文沉淀)
 ├─ graph_nodes           (图谱节点: paper/concept/related)
 │    └─ graph_edges      (图谱边: cites/related/concept_of)
 ├─ ai_conversations      (AI 会话)
 │    └─ ai_messages      (对话消息)
 ├─ import_records        (导入记录)
 ├─ export_records        (Obsidian 导出 / 数据备份)
 ├─ teams                (团队，owner=队长)
 │    ├─ team_members     (成员，一人可加入多团队)
 │    └─ team_applications (加入申请，队长审批)
 ├─ paper_shares         (论文分享: user/team，可撤销)
 ├─ friendships          (好友，分享互动自动沉淀)
 └─ user_preferences     (偏好/引擎配置/语言)
''')
# ================= 4 表设计 =================
h1('4. 数据表详细设计')

def table_block(name, comment, cols):
    h2(f'4.{name}')
    para(comment, size=9.5, color=RGBColor(0x64,0x74,0x8B))
    table(['字段名', '类型', '允许空', '默认', '说明'],
          [[c[0], c[1], c[2], c[3], c[4]] for c in cols],
          widths=[3.6, 3.0, 1.3, 2.2, 5.9])

table_block('1 users 用户表', '注册/登录/个人资料；头像支持 emoji 与图片；研究方向以顿号分隔多值存储。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['username', 'VARCHAR(50)', 'NO', '-', '昵称（可修改）'],
    ['email', 'VARCHAR(100)', 'NO', '-', '登录账号，UNIQUE'],
    ['password_hash', 'VARCHAR(255)', 'NO', '-', 'bcrypt/argon2 哈希'],
    ['id_card', 'VARCHAR(18)', 'YES', 'NULL', '身份证（脱敏展示用）'],
    ['role', 'VARCHAR(20)', 'NO', '普通用户', '角色：普通用户/技术负责人/团队管理员/专业版'],
    ['avatar', 'VARCHAR(255)', 'NO', '🦉', '头像：emoji 或图片 URL/base64'],
    ['research_direction', 'VARCHAR(255)', 'YES', 'NULL', '研究方向，顿号分隔（用于按维度推荐论文）'],
    ['status', 'TINYINT', 'NO', '1', '0=禁用 1=正常'],
    ['deleted_at', 'DATETIME', 'YES', 'NULL', '软删除'],
    ['created_at / updated_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '创建/更新时间'],
])

table_block('2 user_preferences 用户偏好表', '一对一；界面语言、翻译/OCR 引擎配置（原型个人中心「偏好设置」）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users，UNIQUE'],
    ['lang', 'VARCHAR(10)', 'NO', 'zh', '界面语言：zh/en（中英切换）'],
    ['translate_engine', 'VARCHAR(50)', 'YES', 'NULL', '翻译引擎标识：内置/自定义服务商'],
    ['translate_config', 'JSON', 'YES', 'NULL', '翻译引擎配置：{provider,url,api_key,model,timeout}'],
    ['ocr_engine', 'VARCHAR(50)', 'YES', 'NULL', 'OCR 引擎标识'],
    ['ocr_config', 'JSON', 'YES', 'NULL', 'OCR 引擎配置：{provider,url,api_key,model,timeout}'],
    ['default_category', 'BIGINT UNSIGNED', 'YES', 'NULL', '默认导入类别（FK→categories）'],
    ['created_at / updated_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('3 categories 类别表', '论文大类别（文件夹），用户级；原型「我的文献」左侧分类。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['name', 'VARCHAR(50)', 'NO', '-', '类别名，UNIQUE(user_id,name)'],
    ['sort', 'INT', 'NO', '0', '排序'],
    ['is_system', 'TINYINT', 'NO', '0', '系统预设类别'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('4 papers 论文表', '个人文献库主表；元数据 + 阅读状态 + 进度 + 来源（本地文件/DOI/arXiv/收藏网站）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users（个人文献库）'],
    ['category_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→categories'],
    ['title', 'VARCHAR(500)', 'NO', '-', '英文标题'],
    ['title_zh', 'VARCHAR(500)', 'YES', 'NULL', '中文标题'],
    ['authors', 'VARCHAR(500)', 'YES', 'NULL', '作者列表'],
    ['venue', 'VARCHAR(200)', 'YES', 'NULL', '期刊/会议'],
    ['year', 'SMALLINT', 'YES', 'NULL', '发表年份'],
    ['doi', 'VARCHAR(100)', 'YES', 'NULL', 'DOI，UNIQUE(user_id,doi)'],
    ['arxiv_id', 'VARCHAR(50)', 'YES', 'NULL', 'arXiv ID'],
    ['abstract', 'MEDIUMTEXT', 'YES', 'NULL', '英文摘要'],
    ['abstract_zh', 'MEDIUMTEXT', 'YES', 'NULL', '中文摘要'],
    ['intro', 'VARCHAR(1000)', 'YES', 'NULL', '一句话简介（可 AI 生成/编辑）'],
    ['tags', 'VARCHAR(255)', 'YES', 'NULL', '逗号分隔标签'],
    ['status', 'VARCHAR(20)', 'NO', '想读', '想读/在读/读完'],
    ['starred', 'TINYINT', 'NO', '0', '⭐ 收藏'],
    ['read_progress', 'TINYINT UNSIGNED', 'NO', '0', '阅读进度 0-100'],
    ['source_type', 'VARCHAR(20)', 'NO', 'file', '来源：file/doi/arxiv/site'],
    ['file_type', 'VARCHAR(10)', 'YES', 'NULL', '主文件格式：pdf/docx/pptx/txt'],
    ['file_path', 'VARCHAR(500)', 'YES', 'NULL', '本地文件路径'],
    ['file_size', 'BIGINT', 'YES', 'NULL', '文件大小（字节）'],
    ['cover_url', 'VARCHAR(500)', 'YES', 'NULL', '封面图'],
    ['last_read_at', 'DATETIME', 'YES', 'NULL', '最近阅读时间'],
    ['deleted_at', 'DATETIME', 'YES', 'NULL', '软删除'],
    ['created_at / updated_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('5 paper_files 论文文件表', '一篇论文多格式附件（PDF/Word/PPT），支持多格式切换阅读。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['paper_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→papers'],
    ['file_type', 'VARCHAR(10)', 'NO', '-', 'pdf/word/ppt/txt'],
    ['file_name', 'VARCHAR(255)', 'NO', '-', '原始文件名'],
    ['file_path', 'VARCHAR(500)', 'NO', '-', '存储路径'],
    ['file_size', 'BIGINT', 'YES', 'NULL', '字节数'],
    ['page_count', 'INT', 'YES', 'NULL', '页数/幻灯片数'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('6 paper_highlights 高亮表', '阅读器划词高亮（颜色 y/g/b/p/o）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→papers'],
    ['para_index', 'INT', 'NO', '-', '所在段落序号（双栏 data-i）'],
    ['sel_text', 'VARCHAR(1000)', 'NO', '-', '划选原文'],
    ['color', 'VARCHAR(10)', 'NO', 'y', '高亮颜色'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('7 paper_notes 笔记表', '划词笔记与截图 OCR 笔记统一存储；AI 翻译/总结可人工修改后保存。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→papers'],
    ['highlight_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→paper_highlights（划词笔记关联高亮）'],
    ['sel_text', 'VARCHAR(1000)', 'YES', 'NULL', '划选/截图对应文本'],
    ['note_text', 'TEXT', 'YES', 'NULL', '笔记正文'],
    ['ai_translation', 'TEXT', 'YES', 'NULL', 'AI 翻译（用户可修改后保存）'],
    ['ai_summary', 'TEXT', 'YES', 'NULL', 'AI 总结（用户可修改后保存）'],
    ['source', 'VARCHAR(20)', 'NO', 'selection', '来源：selection 划词 / ocr 截图'],
    ['ocr_image_path', 'VARCHAR(500)', 'YES', 'NULL', '截图图片路径（source=ocr）'],
    ['ocr_rect', 'VARCHAR(50)', 'YES', 'NULL', '截图区域（如 300×120 px）'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('8 glossary_terms 术语库', '跨论文去重沉淀的术语（阅读器划词翻译命中、知识库术语列表）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['term_en', 'VARCHAR(200)', 'NO', '-', '英文术语，UNIQUE(user_id,term_en)'],
    ['term_zh', 'VARCHAR(200)', 'NO', '-', '中文译名'],
    ['description', 'VARCHAR(500)', 'YES', 'NULL', '说明'],
    ['source_paper_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→papers 来源论文'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('9 user_sources 收藏网站表', '用户级收藏网站；新用户注册时预置 6 个默认来源（is_default=1）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['name', 'VARCHAR(50)', 'NO', '-', '站点名（arXiv/Semantic Scholar 等）'],
    ['url', 'VARCHAR(500)', 'NO', '-', '站点地址'],
    ['icon', 'VARCHAR(200)', 'YES', 'NULL', '站点图标'],
    ['source_type', 'VARCHAR(20)', 'NO', 'builtin', 'builtin 预置 / custom 自定义'],
    ['is_default', 'TINYINT', 'NO', '0', '是否为默认预置来源（新用户初始化标记）'],
    ['sort', 'INT', 'NO', '0', '排序'],
    ['enabled', 'TINYINT', 'NO', '1', '启用状态'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('10 graph_nodes 图谱节点表', '知识图谱节点：库内论文(paper)、概念(concept)、外部相关论文(related)。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['node_type', 'VARCHAR(20)', 'NO', '-', 'paper/concept/related'],
    ['label', 'VARCHAR(300)', 'NO', '-', '节点显示名（论文标题/概念名）'],
    ['year', 'SMALLINT', 'YES', 'NULL', '年份（论文节点）'],
    ['cites', 'INT', 'YES', 'NULL', '被引量'],
    ['tags', 'VARCHAR(255)', 'YES', 'NULL', '标签'],
    ['paper_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→papers（node_type=paper 关联库内论文）'],
    ['read_status', 'TINYINT', 'NO', '0', '是否已入库/已读（相关论文点「导入」后置 1）'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('11 graph_edges 图谱边表', '节点关系边：引用/关联/概念归属；点击节点推荐相关论文基于此表。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['source_node_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→graph_nodes 起点'],
    ['target_node_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→graph_nodes 终点'],
    ['relation_type', 'VARCHAR(20)', 'NO', '-', 'cites 引用 / related 相关 / concept_of 概念归属'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('12 ocr_records 截图 OCR 记录表', '阅读器截图 OCR 全过程留痕：截图 → 识别 → AI 翻译总结 → 人工修改。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→papers'],
    ['image_path', 'VARCHAR(500)', 'NO', '-', '截图图片路径'],
    ['rect', 'VARCHAR(50)', 'YES', 'NULL', '截图区域尺寸'],
    ['ocr_text', 'TEXT', 'NO', '-', 'OCR 识别结果（可编辑）'],
    ['ai_translation', 'TEXT', 'YES', 'NULL', 'AI 翻译（可编辑）'],
    ['ai_summary', 'TEXT', 'YES', 'NULL', 'AI 总结（可编辑）'],
    ['is_edited', 'TINYINT', 'NO', '0', '是否被人工修改过'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('13 reading_records 阅读记录表', '每篇论文聚合一条：进度、累计时长、最近章节；支撑「继续阅读」与统计。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→papers，UNIQUE(user_id,paper_id)'],
    ['progress', 'TINYINT UNSIGNED', 'NO', '0', '进度 0-100'],
    ['duration_sec', 'INT UNSIGNED', 'NO', '0', '累计精读时长（秒）'],
    ['last_section', 'VARCHAR(50)', 'YES', 'NULL', '最近章节（大纲 id）'],
    ['last_position', 'INT', 'YES', 'NULL', '滚动位置'],
    ['read_at', 'DATETIME', 'YES', 'NULL', '最近阅读时间'],
    ['created_at / updated_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('14 ai_conversations AI 会话表', '阅读器右侧 AI 对话会话（可绑定论文上下文）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→papers 会话关联论文（可为空）'],
    ['title', 'VARCHAR(200)', 'YES', 'NULL', '会话标题'],
    ['created_at / updated_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('15 ai_messages AI 消息表', '会话内消息（user/assistant）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['conversation_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→ai_conversations'],
    ['role', 'VARCHAR(10)', 'NO', '-', 'user / assistant'],
    ['content', 'TEXT', 'NO', '-', '消息内容'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('16 import_records 导入记录表', '文献导入流水：本地文件（单个/批量）、收藏网站、DOI/arXiv。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→papers 导入成功后的论文'],
    ['file_name', 'VARCHAR(255)', 'YES', 'NULL', '原文件名'],
    ['file_path', 'VARCHAR(500)', 'YES', 'NULL', '文件路径'],
    ['import_type', 'VARCHAR(20)', 'NO', '-', 'file 本地文件 / site 收藏网站 / doi / arxiv'],
    ['status', 'VARCHAR(20)', 'NO', 'pending', 'pending/success/failed'],
    ['error_msg', 'VARCHAR(500)', 'YES', 'NULL', '失败原因'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

table_block('17 export_records Obsidian 导出记录表', '知识库导出 Obsidian 历史（每篇论文一个 .md）。', [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['paper_ids', 'JSON', 'YES', 'NULL', '导出的论文 id 清单'],
    ['file_count', 'INT', 'NO', '0', '导出文件数'],
    ['export_type', 'VARCHAR(20)', 'NO', 'multi', 'single 单篇 / multi 批量 / zip 打包'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
])

# ================= 5 索引 =================
h1('5. 索引设计')
table(['表', '索引', '类型', '说明'], [
    ['users', 'uk_email', 'UNIQUE', '登录账号唯一'],
    ['papers', 'idx_papers_user_status', 'NORMAL', '用户 + 状态筛选'],
    ['papers', 'idx_papers_user_cat', 'NORMAL', '类别文件夹'],
    ['papers', 'idx_papers_user_star', 'NORMAL', '⭐ 收藏列表'],
    ['papers', 'uk_papers_user_doi', 'UNIQUE', '同用户 DOI 去重'],
    ['paper_notes', 'idx_notes_user_paper', 'NORMAL', '笔记聚合查询'],
    ['glossary_terms', 'uk_terms_user_en', 'UNIQUE', '术语去重'],
    ['user_sources', 'idx_sources_user', 'NORMAL', '用户收藏网站'],
    ['graph_edges', 'idx_edges_user_src', 'NORMAL', '图谱邻居查询（推荐）'],
    ['graph_edges', 'idx_edges_user_tgt', 'NORMAL', '反向邻居'],
    ['reading_records', 'uk_read_user_paper', 'UNIQUE', '每用户每论文一条'],
    ['reading_records', 'idx_read_user_at', 'NORMAL', '最近阅读排序'],
    ['ai_messages', 'idx_msg_conv', 'NORMAL', '会话消息'],
    ['import_records', 'idx_import_user_status', 'NORMAL', '导入记录查询'],
], widths=[3.4, 4.6, 1.8, 6.2])

# ================= 6 业务 SQL =================
h1('6. 关键业务流程 SQL 示例')

h2('6.1 新用户注册初始化（预置 6 个默认收藏网站）')
code_block('''-- 事务：建用户 + 偏好 + 默认来源 + 默认类别
START TRANSACTION;
INSERT INTO users (username, email, password_hash, role, avatar, research_direction)
VALUES ('王工', 'wang@xx.com', '<bcrypt>', '普通用户', '🦉', 'AI视觉识别、OCR、考试防作弊');
SET @uid = LAST_INSERT_ID();

INSERT INTO user_preferences (user_id, lang, translate_engine, ocr_engine)
VALUES (@uid, 'zh', 'builtin', 'builtin');

INSERT INTO user_sources (user_id, name, url, source_type, is_default, sort) VALUES
(@uid, 'arXiv', 'https://arxiv.org', 'builtin', 1, 1),
(@uid, 'Semantic Scholar', 'https://www.semanticscholar.org', 'builtin', 1, 2),
(@uid, 'ACL Anthology', 'https://aclanthology.org', 'builtin', 1, 3),
(@uid, 'OpenReview', 'https://openreview.net', 'builtin', 1, 4),
(@uid, '百度学术', 'https://xueshu.baidu.com', 'builtin', 1, 5),
(@uid, '谷歌学术', 'https://scholar.google.com', 'builtin', 1, 6);

INSERT INTO categories (user_id, name, is_system) VALUES
(@uid,'Transformer',1),(@uid,'OCR',1),(@uid,'目标检测',1),(@uid,'图像分类',1);
COMMIT;''')

h2('6.2 本地导入文献（文件 + 表单：类别/状态/标签/简介）')
code_block('''INSERT INTO import_records (user_id, file_name, file_path, import_type, status)
VALUES (@uid, 'paper.pdf', '/files/u1/p4.pdf', 'file', 'pending');

INSERT INTO papers (user_id, category_id, title, authors, venue, year, intro,
                    tags, status, source_type, file_type, file_path, file_size)
VALUES (@uid, 2, 'PP-OCR: A Practical Ultra Lightweight OCR System', 'Yuning Du, et al.',
        'arXiv', 2020, '轻量级实用 OCR 系统，兼顾精度与速度。', 'OCR,轻量,端到端',
        '想读', 'file', 'pdf', '/files/u1/p4.pdf', 5242880);
SET @pid = LAST_INSERT_ID();

INSERT INTO paper_files (paper_id, file_type, file_name, file_path, file_size) VALUES
(@pid, 'pdf', 'paper.pdf', '/files/u1/p4.pdf', 5242880),
(@pid, 'ppt', 'slides.pptx', '/files/u1/p4-ppt.pptx', 2097152);   -- 多格式附件

UPDATE import_records SET paper_id = @pid, status = 'success' WHERE id = LAST_INSERT_ID();''')

h2('6.3 阅读进度上报（阅读器滚动，节流保存）')
code_block('''INSERT INTO reading_records (user_id, paper_id, progress, duration_sec, last_section, last_position, read_at)
VALUES (@uid, @pid, 42, 360, 's4', 1782, NOW())
ON DUPLICATE KEY UPDATE
  progress = VALUES(progress),
  duration_sec = duration_sec + VALUES(duration_sec),
  last_section = VALUES(last_section),
  last_position = VALUES(last_position),
  read_at = NOW();

-- 「继续阅读」查询
SELECT p.*, r.progress, r.last_section
FROM papers p LEFT JOIN reading_records r ON r.paper_id = p.id AND r.user_id = @uid
WHERE p.user_id = @uid AND p.deleted_at IS NULL
ORDER BY r.read_at DESC LIMIT 10;''')

h2('6.4 划词笔记（高亮 + 笔记 + AI 翻译/总结，人工修改后保存）')
code_block('''INSERT INTO paper_highlights (user_id, paper_id, para_index, sel_text, color)
VALUES (@uid, @pid, 3, 'The transformer is the core...', 'y');
SET @hid = LAST_INSERT_ID();

INSERT INTO paper_notes (user_id, paper_id, highlight_id, sel_text, note_text,
                         ai_translation, ai_summary, source)
VALUES (@uid, @pid, @hid, 'The transformer is the core...', '重点句',
        'Transformer 是本论文核心，引入自注意力机制。',   -- 人工修改后保存
        '介绍核心概念，为方法阐述铺垫。', 'selection');

-- 知识库「我的笔记」聚合
SELECT p.id, p.title, COUNT(n.id) AS note_count, GROUP_CONCAT(n.note_text SEPARATOR '|') AS notes
FROM paper_notes n JOIN papers p ON p.id = n.paper_id
WHERE n.user_id = @uid GROUP BY p.id, p.title;''')

h2('6.5 截图 OCR + AI 翻译总结（三块可编辑，存为笔记）')
code_block('''INSERT INTO ocr_records (user_id, paper_id, image_path, rect, ocr_text,
                             ai_translation, ai_summary, is_edited)
VALUES (@uid, @pid, '/ocr/u1/20260816_0930.png', '300×120 px',
        '【人工修正】The network architecture is revised...',  -- 用户编辑后的 OCR 文本
        '【人工译文】该网络架构经人工审校修正。',
        '【人工总结】核对了方法部分，确认了模型结构创新。', 1);

-- 若用户点「存为笔记」，同步生成笔记（关联截图）
INSERT INTO paper_notes (user_id, paper_id, sel_text, note_text, ai_translation,
                         ai_summary, source, ocr_image_path, ocr_rect)
SELECT user_id, paper_id, ocr_text, '截图OCR笔记', ai_translation, ai_summary,
       'ocr', image_path, rect FROM ocr_records WHERE id = LAST_INSERT_ID();''')

h2('6.6 图谱点击节点 → 推荐相关论文（直接/间接，阅读或导入）')
code_block('''-- 一度邻居（直接相关）
SELECT gn.id, gn.label, gn.year, gn.cites, gn.tags, gn.read_status,
       p.id AS paper_id, p.status AS paper_status
FROM graph_edges ge
JOIN graph_nodes gn ON gn.id = IF(ge.source_node_id = @nid, ge.target_node_id, ge.source_node_id)
LEFT JOIN papers p ON p.id = gn.paper_id AND p.user_id = @uid
WHERE (ge.source_node_id = @nid OR ge.target_node_id = @nid)
  AND ge.user_id = @uid AND gn.node_type IN ('paper','related')
  AND gn.id <> @nid;

-- 二度关联（经概念节点）：与一度邻居共享概念节点的其他论文
SELECT gn2.label, gn2.year, gn2.cites, gn2.read_status
FROM graph_edges ge1                       -- 节点-概念
JOIN graph_edges ge2 ON ge2.user_id = @uid -- 概念-其他论文
  AND (ge2.source_node_id = ge1.source_node_id OR ge2.target_node_id = ge1.target_node_id)
JOIN graph_nodes gc ON gc.id = ge1.source_node_id AND gc.node_type = 'concept'
JOIN graph_nodes gn2 ON gn2.id = IF(ge2.source_node_id = gc.id, ge2.target_node_id, ge2.source_node_id)
WHERE ge1.source_node_id = @nid AND gn2.node_type IN ('paper','related')
  AND gn2.id <> @nid AND gn2.id NOT IN (SELECT id FROM graph_nodes WHERE id IN (SELECT paper_id FROM papers WHERE user_id=@uid));

-- 点击「导入」：外部 related 节点入库
UPDATE graph_nodes SET read_status = 1 WHERE id = @rid;
INSERT INTO papers (user_id, title, year, source_type) SELECT @uid, label, year, 'graph' FROM graph_nodes WHERE id = @rid;''')

h2('6.7 导出 Obsidian（每篇一个 .md，含 frontmatter/AI 总结/摘要/术语/笔记）')
code_block('''SELECT p.title, p.title_zh, p.authors, p.venue, p.year, p.doi,
       p.tags, p.status, p.read_progress, p.intro,
       n.ai_summary, p.abstract, p.abstract_zh,
       GROUP_CONCAT(DISTINCT CONCAT(gt.term_en,'=',gt.term_zh) SEPARATOR ';') AS glossary,
       GROUP_CONCAT(DISTINCT n.note_text SEPARATOR '|') AS notes
FROM papers p
LEFT JOIN paper_notes n  ON n.paper_id = p.id AND n.user_id = @uid
LEFT JOIN glossary_terms gt ON gt.user_id = @uid
  AND gt.source_paper_id = p.id
WHERE p.user_id = @uid AND p.id IN (<勾选的论文 id 列表>)
GROUP BY p.id;

INSERT INTO export_records (user_id, paper_ids, file_count, export_type)
VALUES (@uid, JSON_ARRAY(1,2,3), 3, 'multi');''')

h2('6.8 界面中英切换 / 引擎配置保存（个人中心偏好设置）')
code_block('''UPDATE user_preferences
SET lang = 'en', translate_config = JSON_SET(translate_config,
        '$.provider', 'deepseek', '$.url', 'https://api.deepseek.com',
        '$.model', 'deepseek-chat', '$.timeout', 30),
    ocr_config = JSON_SET(ocr_config,
        '$.provider', 'paddleocr', '$.url', 'http://10.1.7.68:8088')
WHERE user_id = @uid;

-- 学习建议：按研究方向维度推荐（顿号分割 → 匹配标签/类别）
SELECT p.id, p.title, p.cat FROM papers p
WHERE p.user_id = @uid
  AND (p.tags LIKE '%OCR%' OR p.cat LIKE '%OCR%')     -- 每个方向一个查询
ORDER BY p.read_progress ASC, p.starred DESC LIMIT 5;''')

# ================= 7 与原型对照 =================
h1('7. 原型功能 → 表对照说明（V2.0 全功能）')
table(['原型功能', '实现说明'], [
    ['右上角头像/个人中心（头像/角色/研究方向可改）', 'users.avatar/role/research_direction，修改走 UPDATE users'],
    ['收藏网站（新用户默认 6 个，用户级隔离 + arXiv 真实检索）', 'user_sources.is_default=1 预置；arXiv 检索走后端代理（原型 arxiv_proxy.py）'],
    ['导入文献 Tab（文件上传单个/批量 + 资料表单）', 'import_records + papers + paper_files；批量即一次事务多行'],
    ['回收站（删除→恢复→彻底删除）', 'papers.deleted_at 软删除；彻底删除=物理 DELETE + 级联清理 notes/highlights'],
    ['阅读器多格式（PDF/Word/PPT）+ 缩放翻页 + 大纲跳转', 'paper_files.file_type 多行；原型模拟渲染，正式版接 PDF.js/Office Online'],
    ['划词自动 AI 翻译+总结（可编辑）→ 存为笔记', 'paper_highlights + paper_notes.ai_translation/ai_summary'],
    ['笔记三档可见性（公开/仅好友/私密，默认公开）', 'paper_notes.visibility；社区动态流仅查 visibility=public'],
    ['截图 OCR + AI 翻译总结（可编辑）', 'ocr_records 全过程留痕 + paper_notes(source=ocr)'],
    ['知识图谱节点点击 → 推荐（阅读/导入）', 'graph_edges 一度/二度查询 + read_status 导入标记'],
    ['知识库（卡片/术语/笔记聚合/统计卡跳转）', 'papers + glossary_terms + paper_notes + reading_records 聚合'],
    ['导出 Obsidian 逐篇 md', 'export_records 记录 + 查询组装（frontmatter/总结/摘要/术语/笔记）'],
    ['数据备份（个人中心导出/导入 JSON）', 'export_records.backup_type 扩展；全量 JSON 序列化/反序列化'],
    ['偏好设置（翻译/OCR 引擎、中英切换、测试连接真实探测）', 'user_preferences.lang/translate_config/ocr_config（JSON）'],
    ['AI 对话 / 问 AI', 'ai_conversations + ai_messages（按会话/论文上下文）'],
    ['社区交流（动态流/论文聚合/点赞评论/分页）', 'paper_notes(public) + 点赞评论表（paper_note_likes/comments 扩展）+ 分页'],
    ['团队（列表/创建/申请/审批/成员管理）', 'teams + team_members + team_applications（队长=owner）'],
    ['论文分享（个人/好友/团队 + 分享管理）', 'paper_shares(target_type) + friendships（分享自动沉淀好友）'],
], widths=[5.5, 10.5])

# ================= 8 团队协作扩展 =================
h1('8. 团队协作扩展设计（个人 ↔ 团队双模式）')
para('产品定位：既是个人精读工具，也可作为团队共读工具。本设计预留共享维度，正式版据此扩展，个人模式（默认）与团队模式共存、互不干扰。', size=10)

h2('8.1 新增表：teams / team_members')
table(['字段名', '类型', '允许空', '默认', '说明'], [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['name', 'VARCHAR(50)', 'NO', '-', '团队名'],
    ['owner_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users 创建者'],
    ['avatar', 'VARCHAR(255)', 'YES', 'NULL', '团队头像'],
    ['invite_code', 'VARCHAR(20)', 'YES', 'NULL', '邀请码（成员加入）'],
    ['created_at / updated_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
], widths=[3.6, 3.0, 1.3, 2.2, 5.9])
table(['字段名', '类型', '允许空', '默认', '说明'], [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['team_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→teams，UNIQUE(team_id,user_id)'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users 成员'],
    ['role', 'VARCHAR(20)', 'NO', 'member', 'owner/admin/member'],
    ['joined_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
], widths=[3.6, 3.0, 1.3, 2.2, 5.9])

h2('8.2 既有表扩展（共享维度）')
table(['表', '新增字段', '说明'], [
    ['papers', 'share_type VARCHAR(20) DEFAULT \'private\'', 'private 个人 / team 团队共享 / public 公开'],
    ['papers', 'team_id BIGINT NULL', 'FK→teams（share_type=team 时归属团队）'],
    ['user_sources', 'team_id BIGINT NULL', '团队共享收藏网站'],
    ['glossary_terms', 'team_id BIGINT NULL', '团队共享术语（去重维度变为 team_id+term_en）'],
    ['paper_notes', 'team_id BIGINT NULL', '团队共享笔记'],
    ['graph_nodes / graph_edges', 'team_id BIGINT NULL', '团队共享图谱'],
    ['categories', 'team_id BIGINT NULL', '团队共享类别'],
], widths=[3.6, 5.4, 7.0])

h2('8.3 双模式数据隔离规则')
code_block('''-- 个人模式（默认）：user_id 归属，team_id 为 NULL —— 现有逻辑不变
SELECT * FROM papers WHERE user_id = @uid AND team_id IS NULL;

-- 团队模式：user_id=本人 或 team_id 属于我所在的团队
SELECT p.* FROM papers p
JOIN team_members tm ON tm.team_id = p.team_id AND tm.user_id = @uid
WHERE p.share_type = 'team';

-- 共享导入：团队论文入库时归属团队（owner 或 admin 可管理）
INSERT INTO papers (user_id, team_id, share_type, title, ...)
VALUES (@uid, @tid, 'team', '...', ...);''')

h2('8.4 团队知识沉淀闭环')
para('团队模式下，成员精读产生的术语自动归入团队术语库（去重键 team_id+term_en）；AI 知识卡片、笔记聚合、知识图谱按团队维度汇总；个人中心「研究方向推荐」在团队模式下可切换为「团队共读推荐」（按团队近期阅读热榜）。导出 Obsidian 支持团队知识库整体导出。', size=10)

# ================= 9 社区交流与分享扩展 =================
h1('9. 社区交流与分享扩展设计')
para('对应原型「社区交流」Tab：公开笔记动态流、论文社区聚合、团队创建/申请/审批、论文分享（个人/好友/团队）。', size=10)

h2('9.1 笔记可见性（三档）')
table(['字段', '说明'], [
    ['paper_notes.visibility', 'VARCHAR(10) DEFAULT \'public\'：public 公开（社区动态流可见）/ friends 仅好友 / private 私密（仅本人）'],
    ['社区动态流查询', 'SELECT * FROM paper_notes WHERE visibility=\'public\'（他人视角）；私密/仅好友笔记按 visibility + friendships 过滤'],
    ['论文社区聚合页', '某论文下全部 public 笔记 + 本人全部笔记（含 private，标注🔒）'],
], widths=[4.5, 11.5])

h2('9.2 新增表：team_applications / paper_shares / friendships')
table(['字段名', '类型', '允许空', '默认', '说明'], [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['team_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→teams'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users 申请人'],
    ['status', 'VARCHAR(20)', 'NO', 'pending', 'pending / approved / rejected'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
], widths=[3.6, 3.0, 1.3, 2.2, 5.9], head_bg='1E293B')
para('（team_applications：加入申请，队长审批）', size=9, color=RGBColor(0x64,0x74,0x8B))
table(['字段名', '类型', '允许空', '默认', '说明'], [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users 分享者'],
    ['paper_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→papers'],
    ['target_type', 'VARCHAR(10)', 'NO', '-', 'user 个人/好友 / team 团队'],
    ['target_user_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→users（target_type=user）'],
    ['target_team_id', 'BIGINT UNSIGNED', 'YES', 'NULL', 'FK→teams（target_type=team）'],
    ['status', 'VARCHAR(20)', 'NO', 'active', 'active / revoked（可撤销）'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
], widths=[3.6, 3.0, 1.3, 2.2, 5.9], head_bg='1E293B')
para('（paper_shares：分享记录；分享给个人时自动写入 friendships 沉淀好友）', size=9, color=RGBColor(0x64,0x74,0x8B))
table(['字段名', '类型', '允许空', '默认', '说明'], [
    ['id', 'BIGINT UNSIGNED', 'NO', '自增', '主键'],
    ['user_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users'],
    ['friend_id', 'BIGINT UNSIGNED', 'NO', '-', 'FK→users 好友，UNIQUE(user_id,friend_id)'],
    ['created_at', 'DATETIME', 'NO', 'CURRENT_TIMESTAMP', '-'],
], widths=[3.6, 3.0, 1.3, 2.2, 5.9], head_bg='1E293B')

h2('9.3 社区/分享关键查询')
code_block('''-- 社区动态流（他人视角：只见 public 笔记，按时间倒序）
SELECT n.*, p.title AS paper_title, u.username, u.avatar
FROM paper_notes n
JOIN papers p ON p.id = n.paper_id AND p.share_type IN ('private','team')  -- 论文可读
JOIN users u ON u.id = n.user_id
WHERE n.visibility = 'public' AND n.user_id <> @uid
ORDER BY n.created_at DESC LIMIT 50;

-- 论文社区聚合（某篇论文的公开讨论 + 我的全部笔记）
SELECT * FROM paper_notes
WHERE paper_id = @pid AND (visibility='public' OR user_id=@uid)
ORDER BY created_at DESC;

-- 分享给我的论文
SELECT s.*, p.title FROM paper_shares s
JOIN papers p ON p.id = s.paper_id
WHERE (s.target_user_id = @uid OR s.target_team_id IN (SELECT team_id FROM team_members WHERE user_id=@uid))
  AND s.status='active' AND s.user_id <> @uid;

-- 团队动态（团队频道，后续扩展）：团队成员的公开笔记
SELECT n.* FROM paper_notes n
JOIN team_members tm ON tm.user_id = n.user_id
WHERE tm.team_id = @tid AND n.visibility='public';''')

# ================= 10 回收站 / 备份 / 运行时能力 =================
h1('10. 回收站、数据备份与运行时能力')
para('以下能力与既有表结合实现，不新增表：', size=10)

h2('10.1 回收站（软删除体系）')
table(['项', '规则'], [
    ['删除论文', 'UPDATE papers SET deleted_at=NOW() WHERE id=@pid；同一事务内级联软删 notes/highlights/ocr_records'],
    ['恢复', 'UPDATE papers SET deleted_at=NULL；恢复其关联数据（deleted_at 同批置 NULL）'],
    ['彻底删除', '物理 DELETE + 级联清理（paper_files/notes/highlights/ocr_records/reading_records）'],
    ['列表过滤', '所有业务查询统一加 WHERE deleted_at IS NULL'],
    ['原型对应', 'library.html 左侧「🗑 回收站」+ 恢复/彻底删除（原型用 _trashed 标记 + pm_trash 副本）'],
], widths=[3.2, 12.8])

h2('10.2 数据备份 / 恢复')
table(['项', '规则'], [
    ['导出', '全量 JSON：users 配置 + papers/notes/highlights/terms/sources/graph/shares/teams 序列化；export_records 记 backup 记录'],
    ['导入', '校验 app 标识与版本 → 事务写入 → 前端 reload；原型个人中心「💾 数据备份」'],
    ['换环境迁移', '导出文件可在另一浏览器导入，实现数据迁移'],
], widths=[3.2, 12.8])

h2('10.3 分页约定（社区/知识库/文献库统一）')
para('原型统一分页组件：页码折叠 + 每页条数切换（5/10/20/50）。知识库卡片 6/页、术语 10/页、笔记 5 组/页、社区动态 5/页、团队 4/页、分享 5/页、文献 10/页。正式版统一后端 LIMIT/OFFSET + 总数 COUNT 一次返回。', size=10)

h2('10.4 引擎「测试连接」真实探测')
para('偏好设置中翻译/OCR 引擎地址：浏览器 fetch（mode:no-cors）探测可达性并展示真实延迟；正式版由后端发起带 Key 的鉴权探测（HTTPS GET /v1/models 等端点）。', size=10)

h2('10.5 社区点赞与评论（第 9 章补充）')
table(['表', '字段要点'], [
    ['paper_note_likes', 'id / note_id FK / user_id FK / UNIQUE(note_id,user_id) / created_at'],
    ['paper_note_comments', 'id / note_id FK / user_id FK / content VARCHAR(1000) / created_at（楼层式，原型演示数据可预置）'],
], widths=[5.0, 11.0])

doc.save(r'C:\Users\23677\Documents\paper-reader\docs\PaperMind-数据库设计文档.docx')
print('OK 已生成：docs/PaperMind-数据库设计文档.docx')
